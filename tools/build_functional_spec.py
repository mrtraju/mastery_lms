from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Aegis_Mastery_LMS_Functional_Specification.docx"

BLUE = "1F5EFF"
NAVY = "102033"
LIGHT_BLUE = "EAF1FF"
LIGHT_GRAY = "F3F5F8"
MID_GRAY = "D9DEE7"
GREEN = "0F8A5F"
AMBER = "B7791F"
RED = "C2410C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    text = wrap_cell_text(text)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def wrap_cell_text(text: str) -> str:
    if len(text) < 30:
        return text
    wrapped = text
    for sep in ("\\", "/", "_"):
        wrapped = wrapped.replace(sep, sep + " ")
    return wrapped


def set_cell_width(cell, width: float) -> None:
    twips = int(width * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(twips))
    cell.width = Inches(width)


def set_table_borders(table, color: str = MID_GRAY) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_table(table, header_fill: str = NAVY) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.margin_top = Inches(0.05)
            cell.margin_bottom = Inches(0.05)
            cell.margin_left = Inches(0.07)
            cell.margin_right = Inches(0.07)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
            elif row_idx % 2 == 0:
                set_cell_shading(cell, "FAFBFC")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(2)
    heading.paragraph_format.space_after = Pt(3)
    run = heading.add_run(" | ".join(headers))
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    run.font.size = Pt(9)
    for row_data in rows:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(2)
        if len(row_data) == 2:
            label = p.add_run(f"{row_data[0]}: ")
            label.bold = True
            p.add_run(row_data[1])
        else:
            first = p.add_run(str(row_data[0]))
            first.bold = True
            p.add_run(" - ")
            parts = []
            for header, value in zip(headers[1:], row_data[1:]):
                parts.append(f"{header}: {value}")
            p.add_run("; ".join(parts))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return None


def add_bullets(doc: Document, items: list[str], style: str = "List Bullet") -> None:
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc: Document, title: str, text: str, fill: str = LIGHT_BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.16)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    p2.add_run(text)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 28, NAVY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 12.5, BLUE),
        ("Heading 3", 10.5, NAVY),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name in ("Title", "Heading 1") else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    for list_style in ("List Bullet", "List Number"):
        styles[list_style].font.name = "Aptos"
        styles[list_style].font.size = Pt(9.3)
        styles[list_style].paragraph_format.space_after = Pt(2)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Aegis Mastery LMS Functional Specification | Generated from project source | 2026-04-26").font.size = Pt(8)


def cover(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AEGIS MASTERY LMS")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Functional Specification Document")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated from the project source code in C:\\Users\\rajum\\Projects\\Antigravity\\mastery_lms")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 100, 115)

    doc.add_paragraph()
    add_table(
        doc,
        ["Item", "Details"],
        [
            ["Project", "Aegis Mastery LMS"],
            ["Document type", "Functional specification"],
            ["Generated on", "2026-04-26"],
            ["Source reviewed", "FastAPI backend, SQLAlchemy models, Next.js frontend components, API client, seed data, sample upload assets"],
            ["Audience", "Product owners, developers, QA testers, lecturers, administrators, and implementation reviewers"],
        ],
        widths=[1.55, 5.65],
    )

    add_callout(
        doc,
        "Document basis",
        "This specification describes behavior implemented or directly implied by the current project files. Where the code contains placeholders, temporary scripts, or integration assumptions, those are identified as implementation notes rather than completed production behavior.",
        fill="EFF6FF",
    )
    doc.add_page_break()


def document_control(doc: Document) -> None:
    add_heading(doc, "Document Control", 1)
    add_table(
        doc,
        ["Field", "Specification"],
        [
            ["Version", "1.0"],
            ["Prepared for", "Mastery LMS project workspace"],
            ["Prepared by", "Codex"],
            ["Functional baseline", "Source files present on 2026-04-26"],
            ["Primary modules", "Authentication, student learning, lecturer curriculum management, administrator control, adaptive recommendation engine, content uploads, mastery assessment"],
        ],
        widths=[1.6, 5.7],
    )

    add_heading(doc, "Table of Contents", 1)
    add_numbered(
        doc,
        [
            "Executive Summary",
            "System Scope",
            "Stakeholders and User Roles",
            "Functional Architecture",
            "Data Model",
            "Functional Requirements",
            "User Journeys",
            "API Specification",
            "Business Rules and Algorithms",
            "Validation, Error Handling, and Security",
            "Non-Functional Requirements",
            "Test Scenarios",
            "Open Issues and Implementation Notes",
            "Appendix: Source Traceability",
        ],
    )
    doc.add_page_break()


def executive_summary(doc: Document) -> None:
    add_heading(doc, "Executive Summary", 1)
    doc.add_paragraph(
        "Aegis Mastery LMS is an adaptive learning management system for delivering curriculum content through VARK learning modalities: Visual, Auditory, Reading/Writing, and Kinesthetic. The system combines prior academic score, diagnostic performance, engagement telemetry, and mastery attempts to recommend a content modality for each student."
    )
    doc.add_paragraph(
        "The application is implemented as a FastAPI backend with a Next.js frontend. The backend stores users, curriculum objects, parsed Excel curriculum structures, uploaded learning assets, VARK vectors, mastery logs, and configurable expertise weights. The frontend provides distinct workspaces for students, lecturers, and administrators."
    )
    add_callout(
        doc,
        "Core product objective",
        "Help students reach mastery by selecting the learning modality most likely to help them succeed, rotating away from modalities that have recently produced failed attempts, and surfacing intervention signals to lecturers.",
    )
    add_table(
        doc,
        ["Capability", "Summary"],
        [
            ["Adaptive delivery", "Selects the next modality using the student's VARK vector and prior failed modalities."],
            ["Mastery gate", "Marks a topic as mastered when quiz score is at least 90 percent."],
            ["Diagnostics", "Captures baseline performance and derives an initial VARK distribution during student registration."],
            ["Lecturer tooling", "Allows curriculum Excel upload, extracted topic table review, topic-level material upload, and intervention monitoring."],
            ["Admin tooling", "Supports expertise weight calibration, user directory management, password reset, and module creation/removal."],
        ],
        widths=[1.7, 5.55],
    )


def scope(doc: Document) -> None:
    add_heading(doc, "System Scope", 1)
    add_heading(doc, "In Scope", 2)
    add_bullets(
        doc,
        [
            "User registration, login, token storage, and role-based workspace routing.",
            "Student dashboard showing cognitive profile, module list, title selection, topic selection, adaptive lesson playback, and assessment submission.",
            "Lecturer dashboard with class profile, intervention heatmap, module resource controls, Excel curriculum extraction, and topic material uploads.",
            "Administrator console with expertise weight settings, user management, password reset, and curriculum module management.",
            "Backend persistence for users, settings, curriculum, uploaded asset references, quiz questions, and mastery logs.",
            "Static serving of uploaded learning assets through the FastAPI uploads mount.",
        ],
    )
    add_heading(doc, "Out of Scope or Placeholder", 2)
    add_bullets(
        doc,
        [
            "Production AI content generation is represented by placeholder output in ai_service.py. The OpenAI client is initialized if an API key exists, but the live API call is commented out.",
            "Equivalency auditing currently returns true and does not perform a completed semantic audit.",
            "Password recovery UI text exists, but no forgot-password flow is implemented.",
            "Average mastery is displayed as a fixed 82 percent value in the lecturer dashboard rather than calculated from logs.",
            "The standalone test_student.py script references older /api routes and an old lecturer password; it is not aligned with the current backend route surface.",
        ],
    )


def stakeholders(doc: Document) -> None:
    add_heading(doc, "Stakeholders and User Roles", 1)
    add_table(
        doc,
        ["Role", "Primary goals", "Main permissions"],
        [
            ["Student", "Consume assigned modules, complete topics, receive adaptive content, pass mastery assessments.", "View modules, view profile, request next lesson, submit mastery, send telemetry."],
            ["Lecturer", "Monitor class progress and intervention needs, upload curriculum structure and topic learning resources.", "View student profiles, view failed mastery logs, list curriculum, upload Excel and topic materials."],
            ["Administrator", "Configure system rules, manage users, and manage modules.", "All lecturer permissions plus settings, user directory, user deletion, password reset, module creation and deletion."],
        ],
        widths=[1.15, 3.0, 3.1],
    )
    add_heading(doc, "Role Routing", 2)
    add_bullets(
        doc,
        [
            "After login, the frontend fetches /users/me and uses the returned role to render the appropriate workspace.",
            "The backend validates protected endpoints through JWT bearer authentication and check_role role checks.",
            "The registration form allows student, lecturer, and administrator role selection. For students, registration includes diagnostic calibration before account creation.",
        ],
    )


def architecture(doc: Document) -> None:
    add_heading(doc, "Functional Architecture", 1)
    add_table(
        doc,
        ["Layer", "Technology", "Functional responsibility"],
        [
            ["Frontend", "Next.js 16, React 19, TypeScript, Tailwind CSS", "Authentication UI, role-specific dashboards, lesson playback, quiz submission, upload workflows."],
            ["API backend", "FastAPI", "REST endpoints, auth enforcement, curriculum operations, telemetry, recommendations, mastery logging."],
            ["Persistence", "SQLAlchemy with SQLite by default", "Users, weighted settings, curriculum objects, JSON modality payloads, mastery logs."],
            ["File storage", "Local backend/uploads folder", "Excel uploads, topic assets, PDFs, audio files, kinesthetic files, quiz spreadsheets."],
            ["AI service", "OpenAI client placeholder", "Intended transformation of source text into VARK content and equivalency auditing."],
        ],
        widths=[1.2, 2.1, 3.95],
    )
    add_heading(doc, "Runtime Configuration", 2)
    add_table(
        doc,
        ["Setting", "Default or behavior"],
        [
            ["DATABASE_URL", "Defaults to sqlite:///./mastery.db when unset."],
            ["SECRET_KEY", "Defaults to aegeis-adaptive-mastery-secret-key-2026 when unset."],
            ["OPENAI_API_KEY", "Enables OpenAI client initialization, though live generation is not currently called."],
            ["NEXT_PUBLIC_API_URL", "Defaults frontend requests to http://localhost:8000 when unset."],
            ["CORS", "Backend currently allows all origins, methods, and headers."],
        ],
        widths=[2.0, 5.25],
    )


def data_model(doc: Document) -> None:
    add_heading(doc, "Data Model", 1)
    add_table(
        doc,
        ["Entity", "Key fields", "Functional meaning"],
        [
            ["User", "id, email, hashed_password, role, vark_v, vark_a, vark_r, vark_k, pas", "Stores identity, access role, VARK profile, and weighted expertise baseline."],
            ["WeightedSettings", "id, w_prev, w_diag", "Controls how prior academic score and diagnostic score produce entry expertise."],
            ["CurriculumObject", "id, title, code, level, description, excel_path, excel_data_json, quiz_questions, modality_v/a/r/k", "Represents a module and its top-level or topic-level learning content."],
            ["MasteryLog", "id, user_id, curriculum_id, topic_idx, score, modality_used, timestamp", "Captures assessment attempts and supports remediation, progress display, and lecturer intervention monitoring."],
        ],
        widths=[1.45, 2.7, 3.1],
    )
    add_heading(doc, "Important Relationships", 2)
    add_bullets(
        doc,
        [
            "A user can have many mastery logs.",
            "A curriculum object can have many mastery logs.",
            "A mastery log belongs to one user and one curriculum object.",
            "Curriculum Excel extraction stores topic rows as JSON inside CurriculumObject.excel_data_json.",
            "Topic-level materials are stored inside each topic row under materials.v, materials.a, materials.r, materials.k, and materials.quiz.",
        ],
    )


def functional_requirements(doc: Document) -> None:
    add_heading(doc, "Functional Requirements", 1)
    add_heading(doc, "Authentication and Session Management", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Acceptance criteria"],
        [
            ["AUTH-01", "The system shall allow users to register with email, password, role, and previous academic score.", "A new user is stored if the email is unique; duplicate emails return a 400 error."],
            ["AUTH-02", "The system shall hash passwords before persistence.", "Stored user records contain hashed_password rather than plaintext."],
            ["AUTH-03", "The system shall authenticate with email and password and return a bearer token.", "Valid credentials return access_token and token_type bearer; invalid credentials return 401."],
            ["AUTH-04", "The frontend shall store the access token in localStorage under aegis_token.", "Subsequent API calls include Authorization: Bearer token."],
            ["AUTH-05", "The system shall resolve the current authenticated user.", "/users/me returns the user object when the token is valid and null behavior is handled by the frontend when invalid."],
        ],
        widths=[0.75, 3.05, 3.45],
    )

    add_heading(doc, "Student Registration and Calibration", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Acceptance criteria"],
        [
            ["STU-01", "Student registration shall require a multi-modal diagnostic assessment after account info entry.", "Students progress from account info to baseline calibration before final register call."],
            ["STU-02", "The diagnostic assessment shall present selectable V, A, R, and K learning modes for sample topics.", "Each diagnostic step lets the learner select a modality and then mark understanding."],
            ["STU-03", "The system shall calculate diagnostic score from correct or understood responses.", "Score equals correct count divided by total diagnostic questions times 100."],
            ["STU-04", "The system shall derive a VARK vector from selected modalities, time spent, and performance.", "The vector is normalized so the modality values sum to approximately 1.0."],
            ["STU-05", "Entry expertise shall be calculated as E = PAS * Wp + DS * Wd.", "The registration payload stores the calculated expertise value as pas."],
        ],
        widths=[0.75, 3.05, 3.45],
    )

    add_heading(doc, "Student Learning Workspace", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Acceptance criteria"],
        [
            ["LEARN-01", "The student dashboard shall list available curriculum modules.", "The module cards show code, title, description, adaptive badge, and VARK enabled badge."],
            ["LEARN-02", "When curriculum Excel data exists, students shall navigate by title and topic.", "The UI groups rows by forward-filled Title values and exposes topics beneath each title."],
            ["LEARN-03", "The student profile shall display current VARK distribution and primary modality.", "The highest VARK value determines the primary modality shown."],
            ["LEARN-04", "The system shall show mastered topics with a mastered status when progress score is at least 90.", "Topic cards display 100% Mastered when progress[topic_idx] is at least 90."],
            ["LEARN-05", "The lesson player shall request the next recommended lesson for a curriculum and optional topic.", "The API response contains modality, content, quiz, scaffolding flag, and topic title."],
        ],
        widths=[0.75, 3.05, 3.45],
    )

    add_heading(doc, "Adaptive Lesson Delivery", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Acceptance criteria"],
        [
            ["ADAPT-01", "The backend shall select the highest-ranked VARK modality that has not failed for the same curriculum/topic.", "Failed modalities are logs with score below 90 and are excluded when possible."],
            ["ADAPT-02", "If all modalities have failed, the backend shall return the highest VARK modality anyway.", "The recommendation falls back to the top sorted VARK modality."],
            ["ADAPT-03", "Topic-level materials shall override top-level modality content when a topic index and topic materials exist.", "Topic materials v/a/r/k/quiz replace the curriculum-level payload for that topic."],
            ["ADAPT-04", "The response shall flag scaffolding after multiple failed attempts.", "is_scaffolding is true when failed modality count is greater than 2."],
            ["ADAPT-05", "The frontend shall render each modality using modality-appropriate UI.", "Visual/reading use iframe for uploads, auditory uses audio controls, kinesthetic offers a download link for uploaded files."],
        ],
        widths=[0.75, 3.05, 3.45],
    )

    add_heading(doc, "Mastery Assessment", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Acceptance criteria"],
        [
            ["MAST-01", "Each lesson shall present quiz questions associated with the topic or module.", "If no quiz exists, a placeholder self-check question is rendered."],
            ["MAST-02", "Students shall answer all quiz questions before submitting.", "Submit remains disabled until all questions have answers."],
            ["MAST-03", "The frontend shall calculate score from exact matching selected options to answer values.", "Score is rounded correct divided by total questions times 100."],
            ["MAST-04", "The backend shall persist each mastery attempt.", "A MasteryLog record is created with user, curriculum, topic, score, and modality."],
            ["MAST-05", "The system shall treat scores at least 90 as mastered.", "The response returns is_mastered true when score >= 90."],
            ["MAST-06", "Failed attempts shall trigger remediation behavior.", "The frontend alerts the student and reloads the next lesson, causing the recommender to select another modality where available."],
        ],
        widths=[0.75, 3.05, 3.45],
    )

    add_heading(doc, "Lecturer Workspace", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Acceptance criteria"],
        [
            ["LEC-01", "Lecturers shall view enrolled students.", "/lecturer/class-profile returns users with role student."],
            ["LEC-02", "Lecturers shall view intervention status based on failed attempts.", "Students with no fails show NONE, one fail shows REMEDIATION, and two or more fails show CRITICAL."],
            ["LEC-03", "Lecturers shall list curriculum modules.", "Module Resources renders curriculum items and current upload indicators."],
            ["LEC-04", "Lecturers shall upload a master curriculum Excel file for a module.", "The backend saves the file, extracts Chapter, Title, Topic, Topic_Outcome, and Material columns, and stores standardized JSON rows."],
            ["LEC-05", "Lecturers shall review extracted curriculum rows.", "The table modal shows chapter, title, topic, learning outcome, and upload controls."],
            ["LEC-06", "Lecturers shall upload topic-level V, A, R, K, and quiz materials.", "The backend stores files under uploads/material_{module}_{topic}_{type} and updates the relevant topic materials entry."],
            ["LEC-07", "Quiz spreadsheet upload shall parse Question, A, B, C, D, and Answer columns.", "Parsed questions become the topic materials quiz JSON array."],
        ],
        widths=[0.75, 3.05, 3.45],
    )

    add_heading(doc, "Administrator Workspace", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Acceptance criteria"],
        [
            ["ADM-01", "Administrators shall view and update expertise weights.", "The settings page loads w_prev and w_diag and PATCH /admin/settings persists changes."],
            ["ADM-02", "Administrators shall view all users.", "/admin/users returns all user records."],
            ["ADM-03", "Administrators shall delete users.", "DELETE /admin/users/{id} removes the user or returns 404 if missing."],
            ["ADM-04", "Administrators shall reset a user's password.", "PATCH /admin/users/{id}/password hashes and stores the new password."],
            ["ADM-05", "Administrators and lecturers shall create modules.", "POST /admin/curriculum creates a CurriculumObject with pending modality placeholders."],
            ["ADM-06", "Administrators and lecturers shall update or delete modules.", "PATCH and DELETE /admin/curriculum/{id} mutate or remove matching modules with role enforcement."],
        ],
        widths=[0.75, 3.05, 3.45],
    )


def user_journeys(doc: Document) -> None:
    add_heading(doc, "User Journeys", 1)
    add_heading(doc, "Student First-Time Flow", 2)
    add_numbered(
        doc,
        [
            "Open the app and choose registration.",
            "Enter email, password, role as student, and previous academic score.",
            "Complete baseline calibration by choosing modalities and marking understanding.",
            "The frontend calculates diagnostic score, VARK vector, and weighted entry expertise.",
            "The backend creates the account and stores the hashed password, role, pas, and VARK values supplied by the payload.",
            "Student logs in and enters the student workspace.",
            "Student reviews cognitive profile, selects a module, title, and topic.",
            "Lesson player requests a recommended modality and displays content plus quiz.",
            "Student submits the quiz. If score is at least 90, topic returns to the topic list; otherwise the next recommended modality is loaded.",
        ],
    )
    add_heading(doc, "Lecturer Curriculum Upload Flow", 2)
    add_numbered(
        doc,
        [
            "Lecturer logs in and opens the Instructor Console.",
            "Lecturer expands a curriculum module under Module Resources.",
            "Lecturer uploads a master curriculum Excel workbook.",
            "Backend stores the workbook and extracts standardized topic rows.",
            "Lecturer opens the extracted table modal.",
            "Lecturer uploads topic-specific visual, auditory, reading, kinesthetic, or quiz files.",
            "Backend stores each file and updates the topic materials JSON.",
            "Students selecting that topic receive the uploaded topic materials through the adaptive lesson endpoint.",
        ],
    )
    add_heading(doc, "Administrator Control Flow", 2)
    add_numbered(
        doc,
        [
            "Administrator logs in and enters System Core.",
            "Administrator calibrates PAS and diagnostic weights on the Weights tab.",
            "Administrator reviews users, deletes accounts when needed, or opens password reset overlay.",
            "Administrator creates or removes modules on the Modules tab.",
            "Created modules immediately become available through the curriculum list endpoint.",
        ],
    )


def api_spec(doc: Document) -> None:
    add_heading(doc, "API Specification", 1)
    add_table(
        doc,
        ["Method", "Path", "Auth", "Functional purpose"],
        [
            ["GET", "/", "No", "Health/welcome response."],
            ["POST", "/auth/register", "No", "Create a user account."],
            ["POST", "/auth/login", "No", "Authenticate and return JWT bearer token."],
            ["GET", "/users/me", "Bearer", "Return current authenticated user."],
            ["GET", "/admin/settings", "Admin", "Get expertise weight settings; initializes defaults if missing."],
            ["PATCH", "/admin/settings", "Admin", "Update w_prev and w_diag."],
            ["POST", "/student/diagnostic", "No", "Calculate expertise from user PAS and diagnostic score."],
            ["GET", "/student/{user_id}/profile", "No", "Return student profile by id."],
            ["GET", "/student/{user_id}/curriculum/{curriculum_id}/progress", "No", "Return best topic score per topic index."],
            ["POST", "/student/{user_id}/telemetry", "No", "Refine VARK vector based on modality engagement time."],
            ["GET", "/curriculum/{curriculum_id}/next", "No", "Return adaptive lesson payload for user and optional topic."],
            ["POST", "/curriculum/submit", "No", "Persist mastery log and return mastery status."],
            ["POST", "/curriculum/generate", "Lecturer/Admin", "Generate placeholder multimodal content and create curriculum."],
            ["GET", "/curriculum/list", "No", "List curriculum modules."],
            ["GET", "/lecturer/intervention-heatmap", "Lecturer/Admin", "Return all failed mastery logs."],
            ["GET", "/lecturer/class-profile", "Lecturer/Admin", "Return all student users."],
            ["GET", "/admin/users", "Admin", "List all users."],
            ["DELETE", "/admin/users/{user_id}", "Admin", "Delete a user."],
            ["PATCH", "/admin/users/{user_id}/password", "Admin", "Reset user password."],
            ["POST", "/admin/curriculum", "Lecturer/Admin", "Create module with pending modality placeholders."],
            ["PATCH", "/admin/curriculum/{item_id}", "Lecturer/Admin", "Update module metadata."],
            ["DELETE", "/admin/curriculum/{item_id}", "Lecturer/Admin", "Delete module."],
            ["POST", "/lecturer/curriculum/{item_id}/upload", "Lecturer/Admin", "Upload and parse master curriculum Excel."],
            ["POST", "/lecturer/curriculum/{item_id}/assets/{modality}", "Lecturer/Admin", "Upload top-level modality asset."],
            ["POST", "/lecturer/curriculum/{item_id}/quiz", "Lecturer/Admin", "Update top-level quiz questions."],
            ["POST", "/lecturer/curriculum/{item_id}/topic/{topic_idx}/material/{m_type}", "Lecturer/Admin", "Upload topic material or quiz spreadsheet."],
        ],
        widths=[0.7, 2.75, 1.0, 2.85],
    )
    add_callout(
        doc,
        "Security note",
        "Several student and curriculum endpoints are callable without explicit auth dependencies in the current backend implementation, even though the frontend usually sends tokens. See Open Issues for recommended hardening.",
        fill="FFF7ED",
    )


def business_rules(doc: Document) -> None:
    add_heading(doc, "Business Rules and Algorithms", 1)
    add_table(
        doc,
        ["Rule", "Definition", "Current implementation"],
        [
            ["Expertise score", "E = PAS * Wp + DS * Wd", "Implemented in logic.calculate_expertise and mirrored in the registration frontend."],
            ["VARK normalization", "After telemetry adjustment, modality values should sum to 1.0.", "logic.refine_vark increments one modality and divides all values by total."],
            ["Recommendation", "Choose the highest VARK modality that has not failed for that user/curriculum/topic.", "logic.get_recommended_modality sorts VARK values descending and skips failed modalities."],
            ["Mastery threshold", "A score of 90 percent or greater passes the gate.", "logic.check_mastery uses default threshold 90.0."],
            ["Scaffolding signal", "Multiple failures should trigger reduced complexity/remediation support.", "The API sets is_scaffolding when failed modality count is greater than 2; UI displays a warning."],
            ["Intervention status", "Failed mastery logs identify students needing attention.", "Lecturer UI computes NONE, REMEDIATION, or CRITICAL from failed log counts."],
        ],
        widths=[1.45, 2.65, 3.15],
    )


def validation_security(doc: Document) -> None:
    add_heading(doc, "Validation, Error Handling, and Security", 1)
    add_heading(doc, "Validation and Error Handling", 2)
    add_bullets(
        doc,
        [
            "Registration rejects duplicate emails with HTTP 400.",
            "Login rejects invalid credentials with HTTP 401 and WWW-Authenticate bearer header.",
            "Role checks reject unauthorized access with HTTP 403.",
            "Missing users or modules return HTTP 404 in delete/update flows where checks are implemented.",
            "Excel parsing failures return HTTP 400 with a detail message beginning with Excel Failure.",
            "Topic material upload rejects invalid topic indexes with HTTP 400.",
            "Quiz upload parsing failures return HTTP 400 with Quiz Parse Error details.",
        ],
    )
    add_heading(doc, "Security Controls", 2)
    add_bullets(
        doc,
        [
            "Passwords are hashed using passlib CryptContext with bcrypt.",
            "JWT tokens use HS256 and include the user email in the sub claim.",
            "Token expiration is configured for 24 hours.",
            "Protected endpoints use OAuth2 bearer tokens and role checks where dependencies are present.",
            "Uploaded files are stored locally and exposed through the /uploads static route.",
        ],
    )


def non_functional(doc: Document) -> None:
    add_heading(doc, "Non-Functional Requirements", 1)
    add_table(
        doc,
        ["Category", "Requirement"],
        [
            ["Usability", "Role-specific dashboards shall reduce cross-role complexity and show only the workspace relevant to the authenticated user's role."],
            ["Responsiveness", "Frontend layouts shall support single-column mobile views and multi-column desktop views through responsive grid classes."],
            ["Performance", "Core dashboard data should load through parallel requests where applicable, as implemented in AdminPanel and LecturerDashboard."],
            ["Reliability", "Assessment submissions must be persisted before the student receives mastery status."],
            ["Maintainability", "API calls are centralized in frontend/lib/api.ts for consistency."],
            ["Extensibility", "JSON modality payloads and topic materials allow multiple file or generated content types without schema changes."],
            ["Portability", "SQLite is the default database for local development; a non-SQLite DATABASE_URL may be supplied for deployment."],
            ["Accessibility", "Interactive controls should expose labels and keyboard behavior. Current implementation uses native controls for forms, radio inputs, audio, and file inputs, but icon-only controls should receive stronger accessible labels in future hardening."],
        ],
        widths=[1.4, 5.85],
    )


def tests(doc: Document) -> None:
    add_heading(doc, "Test Scenarios", 1)
    add_table(
        doc,
        ["Area", "Scenario", "Expected result"],
        [
            ["Auth", "Register a new student with unique email.", "Account is created after diagnostic completion; user can log in."],
            ["Auth", "Attempt login with wrong password.", "401 response and frontend error alert."],
            ["Admin", "Change PAS and DS weights.", "Settings persist and future registrations use the new weights."],
            ["Student", "Complete a topic quiz with all correct answers.", "Mastery log is saved and is_mastered is true."],
            ["Student", "Fail a quiz in one modality.", "Failed modality is logged; next lesson avoids that modality if another is available."],
            ["Student", "Fail more than two modalities.", "Lesson response has is_scaffolding true and the UI shows a scaffolding notice."],
            ["Lecturer", "Upload curriculum Excel with expected columns.", "Rows are extracted and visible in the table modal."],
            ["Lecturer", "Upload topic quiz spreadsheet.", "Questions parse into JSON and appear in the student's lesson quiz."],
            ["Lecturer", "View intervention heatmap after failures.", "Affected students display REMEDIATION or CRITICAL."],
            ["Admin", "Reset a user's password.", "User can authenticate with the new password."],
        ],
        widths=[1.15, 3.0, 3.1],
    )


def open_issues(doc: Document) -> None:
    add_heading(doc, "Open Issues and Implementation Notes", 1)
    add_table(
        doc,
        ["ID", "Observation", "Recommended action"],
        [
            ["NOTE-01", "The backend UserCreate model does not declare VARK fields, while the registration frontend sends v, a, r, and k. Current backend registration stores role and pas only.", "Add vark_v/a/r/k fields to the request model or map v/a/r/k aliases during registration."],
            ["NOTE-02", "Some student endpoints do not require authentication, including profile, progress, telemetry, next lesson, and mastery submit.", "Require bearer auth and verify the user can access the requested user_id."],
            ["NOTE-03", "CORS allows all origins with credentials enabled.", "Restrict allowed origins per environment."],
            ["NOTE-04", "SECRET_KEY has an in-code default.", "Require environment configuration for deployed environments."],
            ["NOTE-05", "AI generation and equivalency audit are placeholders.", "Implement the actual model call, structured JSON validation, and audit criteria before calling this production-ready."],
            ["NOTE-06", "The lecturer average mastery card is hard-coded at 82 percent.", "Calculate from mastery logs by class, curriculum, or time window."],
            ["NOTE-07", "test_student.py references older routes and outdated credentials.", "Update or replace it with tests against the current API paths."],
            ["NOTE-08", "Uploaded filenames are deterministic by module/topic/type and overwrite prior files.", "Add versioning or unique identifiers if historical materials must be retained."],
            ["NOTE-09", "Role selection is exposed during registration.", "For production, restrict creation of lecturer/admin accounts to admin-invited flows."],
        ],
        widths=[0.85, 3.2, 3.2],
    )


def traceability(doc: Document) -> None:
    add_heading(doc, "Appendix: Source Traceability", 1)
    add_table(
        doc,
        ["Source file", "Specification contribution"],
        [
            ["backend/main.py", "API routes, request/response models, role checks, curriculum upload, recommendation endpoint, mastery submission."],
            ["backend/models.py", "Persistent entities and relationships."],
            ["backend/logic.py", "Expertise, VARK refinement, recommendation, and mastery threshold rules."],
            ["backend/auth.py", "Password hashing, JWT creation, current-user resolution, role enforcement."],
            ["backend/ai_service.py", "AI content generation and equivalency audit placeholder behavior."],
            ["backend/database.py", "Database configuration and SQLite default."],
            ["backend/seed_data.py", "Demo users, default settings, and sample curriculum records."],
            ["frontend/app/page.tsx", "Top-level auth flow, role routing, student module/title/topic navigation."],
            ["frontend/lib/api.ts", "Frontend-to-backend contract and endpoint usage."],
            ["frontend/components/auth/*.tsx", "Login, registration, diagnostic calibration, and VARK survey behavior."],
            ["frontend/components/LessonPlayer.tsx", "Adaptive content rendering and mastery quiz behavior."],
            ["frontend/components/LecturerDashboard.tsx", "Class profile, intervention heatmap, curriculum Excel extraction, topic upload workflows."],
            ["frontend/components/AdminPanel.tsx", "Settings, users, password reset, and module management."],
            ["frontend/components/StudentProfile.tsx", "Student cognitive profile visualization."],
        ],
        widths=[2.35, 4.9],
    )


def build() -> None:
    doc = Document()
    configure_document(doc)
    cover(doc)
    document_control(doc)
    executive_summary(doc)
    scope(doc)
    stakeholders(doc)
    architecture(doc)
    data_model(doc)
    functional_requirements(doc)
    user_journeys(doc)
    api_spec(doc)
    business_rules(doc)
    validation_security(doc)
    non_functional(doc)
    tests(doc)
    open_issues(doc)
    traceability(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
